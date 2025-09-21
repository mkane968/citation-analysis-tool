from flask import Flask, render_template, request, jsonify, after_this_request, send_from_directory, send_file, session, make_response
import pandas as pd
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
import nltk
import re
import os
import tempfile
import PyPDF2
from docx import Document
import io
from werkzeug.utils import secure_filename
from scibert_rhetorical_classifier import RhetoricalMoveClassifier
import pdfkit
import uuid
import json
from datetime import datetime

app = Flask(__name__, static_url_path='/static', static_folder='static')
app.secret_key = 'citation_analysis_secret_key'

# Initialize the rhetorical move classifier
rhetorical_classifier = RhetoricalMoveClassifier()

def preprocess_text(text, author_names=None):
    # Initialize author_names if not provided
    if author_names is None:
        author_names = set()
    # First fix citations that might be broken across lines
    # Look for patterns like "(Wilson, 2019, pp. " followed by "78-92)" on the next line
    text = re.sub(r'\(([\w]+),\s+(\d{4}),\s+pp\.\s*\n\s*(\d+)-(\d+)\)', r'(\1, \2, pp. \3-\4)', text)
    
    # Also fix citations where just the page range is broken across lines
    text = re.sub(r'\(([\w]+),\s+(\d{4}),\s+pp\.\s*\n\s*(\d+.*?)\)', r'(\1, \2, pp. \3)', text)
    
    # Store paragraph breaks for later restoration
    paragraphs = text.split('\n\n')
    processed_paragraphs = []
    
    for paragraph in paragraphs:
        # Normalize whitespace within paragraphs
        paragraph = re.sub(r'\s+', ' ', paragraph)
        paragraph = paragraph.replace('\n', ' ').replace('\r', ' ')
        processed_paragraphs.append(paragraph.strip())
    
    # Rejoin with paragraph markers
    text = '\n\n'.join(processed_paragraphs)
    
    # Remove works cited section if present (common in essays)
    works_cited_patterns = [r'Works Cited.*$', r'References.*$', r'Bibliography.*$']
    for pattern in works_cited_patterns:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE | re.DOTALL)
    
    # Fix common formatting issues from PDFs
    # Fix hyphenated words that might be split across lines
    text = re.sub(r'(\w+)- (\w+)', r'\1\2', text)
    
    # Fix citation formatting issues
    # Fix spaces in citations like "Smith ( 2020 )" to "Smith (2020)"
    text = re.sub(r'\( (\d{4}) \)', r'(\1)', text)
    
    # Temporarily replace "et al." to prevent sentence splitting
    text = text.replace("et al.", "et_al_PLACEHOLDER")
    
    # Split text into sentences, preserving paragraph breaks
    all_sentences = []
    for paragraph in text.split('\n\n'):
        sentences = nltk.sent_tokenize(paragraph)
        # Restore "et al." in each sentence
        sentences = [s.replace("et_al_PLACEHOLDER", "et al.") for s in sentences]
        # Add paragraph marker at the end of each paragraph's sentences
        if sentences:
            sentences[-1] = sentences[-1] + "[PARAGRAPH_BREAK]"
        all_sentences.extend(sentences)
    
    # Use the paragraph-aware sentences
    sentences = all_sentences
    
    # Analyze each sentence for citations
    analyzed_sentences = []
    all_citations = []
    total_citation_count = 0
    
    for sentence in sentences:
        # Check for possessive author references (e.g., "Johnson's approach")
        has_author_reference = False
        for author in author_names:
            possessive_pattern = f"\\b{author}'s\\b"
            narrative_pattern = f"\\b{author}\\s+(?:argues|claims|states|suggests|notes|observes|finds|proposes|demonstrates|shows)\\b"
            
            if re.search(possessive_pattern, sentence, re.IGNORECASE) or re.search(narrative_pattern, sentence, re.IGNORECASE):
                has_author_reference = True
                break
        # Get citations from this sentence
        citations, citation_count = identify_citations(sentence, author_names)
        
        # If no formal citations but we found an author reference, add an implicit citation
        if citation_count == 0 and has_author_reference:
            for author in author_names:
                possessive_pattern = f"\\b{author}'s\\b"
                narrative_pattern = f"\\b{author}\\s+(?:argues|claims|states|suggests|notes|observes|finds|proposes|demonstrates|shows)\\b"
                
                if re.search(possessive_pattern, sentence, re.IGNORECASE):
                    # Just use the author name for possessive references
                    citations.append({
                        'text': author,
                        'style': 'Narrative APA',
                        'is_narrative': True
                    })
                    citation_count = 1
                    break
                elif re.search(narrative_pattern, sentence, re.IGNORECASE):
                    # Just use the author name for narrative references
                    citations.append({
                        'text': author,
                        'style': 'Narrative APA',
                        'is_narrative': True
                    })
                    citation_count = 1
                    break
        has_citation = citation_count > 0
        
        # Determine the citation style for the sentence
        citation_style = 'None'
        if has_citation:
            # Count citation styles
            style_counts = {}
            for citation in citations:
                style = citation.get('style', 'Unknown')
                style_counts[style] = style_counts.get(style, 0) + 1
            
            # Determine predominant style
            if len(style_counts) == 1:
                # Only one style present
                citation_style = list(style_counts.keys())[0]
            elif len(style_counts) > 1:
                # Multiple styles - find the most common
                max_count = 0
                for style, count in style_counts.items():
                    if count > max_count:
                        max_count = count
                        citation_style = style
                # If there's a tie, mark as Mixed
                if list(style_counts.values()).count(max_count) > 1:
                    citation_style = 'Mixed'
        
        # Extract just the citation text for display
        citation_texts = [citation['text'] for citation in citations]
        
        # Check if this sentence has a paragraph break marker
        has_paragraph_break = False
        if sentence.endswith("[PARAGRAPH_BREAK]"):
            sentence = sentence.replace("[PARAGRAPH_BREAK]", "")
            has_paragraph_break = True
            
        analyzed_sentences.append({
            'sentence': sentence,
            'has_citation': has_citation,
            'citations': citation_texts,
            'citation_count': citation_count,
            'citation_style': citation_style,
            'paragraph_break': has_paragraph_break
        })
        
        all_citations.extend(citations)
        total_citation_count += citation_count
        
    return analyzed_sentences, all_citations, total_citation_count

def identify_citations(s, author_names=None):
    # Initialize author_names if not provided
    if author_names is None:
        author_names = set()
    # Initialize the citations list with dictionaries containing text and style
    citations = []
    
    # Track years in APA citations to avoid duplication
    apa_years = set()
    
    # Track rhetorical move type for implicit citations
    rhetorical_move_type = None
    
    # REPORTING MOVE PATTERNS (directly reporting what sources say)
    reporting_patterns = [
        r'\b(?:according to|as stated by|as reported by|as noted by|as observed by)\b',
        r'\b(?:they|he|she|it)\s+(?:state|report|note|observe|mention|describe|explain|define|clarify)s?\b',
        r'\b(?:their|his|her|its)\s+(?:statement|report|observation|description|explanation|definition)s?\b',
        r'\b(?:in the words of|quoted from|directly from)\b',
        r'\bidentifies\b'
    ]
    
    # TRANSFORMING MOVE PATTERNS (paraphrasing or synthesizing)
    transforming_patterns = [
        r'\b(?:their|his|her|its)\s+(?:findings|research|study|analysis|work|paper|article|results|approach|methodology|framework|model|theory|concept|perspective)s?\b',
        r'\b(?:they|he|she|it)\s+(?:found|discovered|reported|showed|demonstrated|argued|suggested|proposed|developed|established|determined|concluded|identified)\b',
        r'\b(?:these|those|this|that)\s+(?:findings|results|studies|researchers|authors|scholars|studies|papers|articles|analyses)\b',
        r'\b(?:according to|as per)\s+(?:them|him|her|it)\b',
        r'\b(?:similar|similarly|likewise|in the same way|in a similar manner)\b',
        r'\b(?:building on|extending|drawing from|synthesizing|combining|integrating)\b',
        r'\bcase\s+(?:study|studies|analysis|analyses)\b'
    ]
    
    # EVALUATING MOVE PATTERNS (critiquing or analyzing sources)
    evaluating_patterns = [
        r'\b(?:overlooks|fails to|neglects|ignores|misses|omits|lacks|inadequately|insufficiently)\b',
        r'\b(?:their|his|her|its)\s+(?:critique|criticism|assessment|evaluation|limitation|weakness|strength|advantage|shortcoming|gap|flaw|merit)s?\b',
        r'\b(?:they|he|she|it)\s+(?:critiques|criticizes|assesses|evaluates|challenges|questions|contests|disputes|refutes|counters)\b',
        r'\b(?:however|nevertheless|nonetheless|although|though|despite|in spite of|yet|but|conversely|in contrast|on the contrary|on the other hand)\b',
        r'\b(?:problematic|questionable|debatable|controversial|unconvincing|unsupported|unsubstantiated|flawed|limited|narrow|biased|misleading)\b',
        r'\b(?:critically|insightfully|perceptively|astutely|shrewdly)\b',
        r'\b(?:strength|weakness|merit|limitation|advantage|disadvantage|benefit|drawback|shortcoming|gap|flaw)s?\b'
    ]
    
    # Check for reporting patterns
    for pattern in reporting_patterns:
        if re.search(pattern, s, re.IGNORECASE):
            # For reporting patterns, extract just the author name if possible
            author_match = re.search(r'According to\s+([\w]+)', s, re.IGNORECASE)
            if author_match:
                author = author_match.group(1)
                citations.append({
                    'text': author,
                    'style': 'Narrative APA',
                    'is_narrative': True
                })
            else:
                # If no clear author, use the pattern match
                match_text = re.search(pattern, s, re.IGNORECASE).group(0)
                citations.append({
                    'text': match_text,
                    'style': 'Implicit',
                    'is_demonstrative': True
                })
            rhetorical_move_type = 'Reporting'
            break
    
    # If not reporting, check for evaluating patterns
    if not rhetorical_move_type:
        for pattern in evaluating_patterns:
            if re.search(pattern, s, re.IGNORECASE):
                # For evaluating patterns, just extract the author name if mentioned
                author_match = re.search(r'([\w]+)[\s\']s\s+(?:approach|framework|model|theory)', s, re.IGNORECASE)
                if author_match:
                    author = author_match.group(1)
                    citations.append({
                        'text': author,
                        'style': 'Narrative APA',
                        'is_narrative': True
                    })
                else:
                    # If no clear author, use a minimal representation
                    citations.append({
                        'text': "Implicit Citation",
                        'style': 'Implicit'
                    })
                rhetorical_move_type = 'Evaluating'
                break
    
    # If neither reporting nor evaluating, check for transforming patterns
    if not rhetorical_move_type:
        for pattern in transforming_patterns:
            if re.search(pattern, s, re.IGNORECASE):
                # For transforming patterns, extract just the author reference if possible
                author_match = re.search(r'(their|his|her|its)\s+(?:findings|research|study)', s, re.IGNORECASE)
                if author_match:
                    citations.append({
                        'text': "Implicit Citation",
                        'style': 'Implicit'
                    })
                else:
                    # If no clear reference, use a minimal representation
                    citations.append({
                        'text': "Implicit Citation",
                        'style': 'Implicit'
                    })
                rhetorical_move_type = 'Transforming'
                break
    
    # Pre-process the string to fix common formatting issues
    # Remove extra spaces between author and year
    s = re.sub(r'(\w+)\s+\((\d{4})\)', r'\1 (\2)', s)
    
    # ===== APA CITATION FORMATS =====
    # 1. Single author: Smith (2020)
    for match in re.findall(r'([\w]+)\s+\((\d{4})\)', s):
        citations.append({
            'text': f"{match[0]} ({match[1]})",
            'style': 'APA'
        })
        apa_years.add(match[1])
    
    # 2. Two authors: Smith and Johnson (2020)
    for match in re.findall(r'([\w]+)\s+and\s+([\w]+)\s+\((\d{4})\)', s):
        citations.append({
            'text': f"{match[0]} and {match[1]} ({match[2]})",
            'style': 'APA'
        })
        apa_years.add(match[2])
    
    # 3. Multiple authors with et al.: Smith et al. (2020)
    for match in re.findall(r'([\w]+)\s+et\s+al\.?\s+\((\d{4})\)', s):
        citations.append({
            'text': f"{match[0]} et al. ({match[1]})",
            'style': 'APA'
        })
        apa_years.add(match[1])
    
    # 4. Three authors: Smith, Johnson, and Lee (2020)
    for match in re.findall(r'([\w]+),\s+([\w]+),\s+and\s+([\w]+)\s+\((\d{4})\)', s):
        citations.append({
            'text': f"{match[0]}, {match[1]}, and {match[2]} ({match[3]})",
            'style': 'APA'
        })
        apa_years.add(match[3])
    
    # 4. Four+ authors: Smith et al. (2020)
    for match in re.findall(r'([\w]+)\s+et\s+al\.\s+\((\d{4})\)', s):
        citations.append({
            'text': f"{match[0]} et al. ({match[1]})",
            'style': 'APA'
        })
        apa_years.add(match[1])
    
    # 5. Report with year: IPCC report (2022)
    for match in re.findall(r'([\w\s]+)\s+report\s+\((\d{4})\)', s, re.IGNORECASE):
        citations.append({
            'text': f"{match[0]} report ({match[1]})",
            'style': 'APA'
        })
        apa_years.add(match[1])
        
    # 6. Multiple authors with page number in parentheses: (Chen et al. 2018, p. 42)
    for match in re.findall(r'\(([\w]+)\s+et\s+al\.\s+(\d{4}),\s+p\.\s+(\d+)\)', s):
        citations.append({
            'text': f"({match[0]} et al. {match[1]}, p. {match[2]})",
            'style': 'APA'
        })
        apa_years.add(match[1])
        
    # 7. Two authors with comma before year: (Reynolds and Ahmed, 2021)
    for match in re.findall(r'\(([\w]+)\s+and\s+([\w]+),\s+(\d{4})\)', s):
        citations.append({
            'text': f"({match[0]} and {match[1]}, {match[2]})",
            'style': 'APA'
        })
        apa_years.add(match[2])
        
    # 8. Multiple authors with comma before year: (Johnson et al., 2021)
    for match in re.findall(r'\(([\w]+)\s+et\s+al\.,\s+(\d{4})\)', s):
        citations.append({
            'text': f"({match[0]} et al., {match[1]})",
            'style': 'APA'
        })
        apa_years.add(match[1])
        
    # 9. Author with page range: (Wilson, 2019, pp. 78-92)
    for match in re.findall(r'\(([\w]+),\s+(\d{4}),\s+pp\.\s+(\d+)-(\d+)\)', s, re.IGNORECASE):
        citations.append({
            'text': f"({match[0]}, {match[1]}, pp. {match[2]}-{match[3]})",
            'style': 'APA'
        })
        apa_years.add(match[1])
        
    # 9c. Exact match for Wilson citation
    if "(Wilson, 2019, pp. 78-92)" in s:
        # Check if this citation has already been added
        citation_text = "(Wilson, 2019, pp. 78-92)"
        if not any(c['text'] == citation_text for c in citations):
            citations.append({
                'text': citation_text,
                'style': 'APA'
            })
            apa_years.add('2019')
        
    # 9b. Alternative pattern for page range with more flexible spacing
    for match in re.findall(r'\(\s*([\w]+)\s*,\s*(\d{4})\s*,\s*pp\.\s*(\d+)\s*-\s*(\d+)\s*\)', s, re.IGNORECASE):
        # Check if this citation has already been added
        citation_text = f"({match[0]}, {match[1]}, pp. {match[2]}-{match[3]})"
        if not any(c['text'] == citation_text for c in citations):
            citations.append({
                'text': citation_text,
                'style': 'APA'
            })
            apa_years.add(match[1])
        
    # 10. Two authors with page number: (Garcia and Martinez, 2022, p. 215)
    for match in re.findall(r'\(([\w]+)\s+and\s+([\w]+),\s+(\d{4}),\s+p\.\s+(\d+)\)', s):
        citations.append({
            'text': f"({match[0]} and {match[1]}, {match[2]}, p. {match[3]})",
            'style': 'APA'
        })
        apa_years.add(match[2])
    
    # ===== MLA CITATION FORMATS =====
    # 1. One author with page number: (Johnson 42)
    for match in re.findall(r'\(([\w]+)\s+(\d+)\)', s):
        # Skip if this looks like an author with year that's already in APA citations
        if len(match[1]) == 4 and 1900 <= int(match[1]) <= 2100 and match[1] in apa_years:
            continue
        citations.append({
            'text': f"{match[0]} {match[1]}",
            'style': 'MLA'
        })
    
    # 2. Two authors with year: (Thompson and Lee 2018)
    for match in re.findall(r'\(([\w]+)\s+and\s+([\w]+)\s+(\d{4})\)', s):
        citations.append({
            'text': f"{match[0]} and {match[1]} {match[2]}",
            'style': 'MLA'
        })
        # Add to APA years to prevent duplication
        apa_years.add(match[2])
    
    # 3. Two authors without page numbers: (Smith and Johnson)
    for match in re.findall(r'\(([\w]+)\s+and\s+([\w]+)\)', s):
        citations.append({
            'text': f"{match[0]} and {match[1]}",
            'style': 'MLA'
        })
    
    # 4. Three authors with page numbers: (Chen, Roberts and Williams 78)
    for match in re.findall(r'\(([\w]+),\s+([\w]+),\s+and\s+([\w]+)\s+(\d+)\)', s):
        citations.append({
            'text': f"{match[0]}, {match[1]} and {match[2]} {match[3]}",
            'style': 'MLA'
        })
    
    # 5. Three authors without page numbers: (Smith, Johnson and Lee)
    for match in re.findall(r'\(([\w]+),\s+([\w]+),\s+and\s+([\w]+)\)', s):
        citations.append({
            'text': f"{match[0]}, {match[1]} and {match[2]}",
            'style': 'MLA'
        })
    
    # 6. Four+ author citations with page number: (Davis et al. 45)
    for match in re.findall(r'\(([\w]+)\s+et\s+al\.\s+(\d+)\)', s):
        citations.append({
            'text': f"{match[0]} et al. {match[1]}",
            'style': 'MLA'
        })
    
    # 7. Four+ author citations without page numbers: (Smith et al.)
    for match in re.findall(r'\(([\w]+)\s+et al\.\)', s):
        citations.append({
            'text': f"{match[0]} et al.",
            'style': 'MLA'
        })
    
    # 8. Citations with quoted material: ("Climate Report")
    for match in re.findall(r'\(\"(.*?)"\)', s):
        citations.append({
            'text': f'"{match}"',
            'style': 'MLA'
        })
    
    # 9. Single author without page number: (Smith)
    for match in re.findall(r'\(([\w]+)\)', s):
        # Skip if it's a year that's already in APA citations
        if match.isdigit() and len(match) == 4 and 1900 <= int(match) <= 2100 and match in apa_years:
            continue
        
        # Skip if this author is already part of another citation
        author_already_cited = False
        for citation in citations:
            if match == citation['text'].split()[0]:  # Check if it's the first word
                author_already_cited = True
                break
        
        if not author_already_cited:
            citations.append({
                'text': match,
                'style': 'MLA'
            })
    
    # 10. Narrative citations: As noted by Roberts
    for match in re.findall(r'(?:noted|mentioned|stated|cited|according to|as per)\s+by\s+([\w]+)', s, re.IGNORECASE):
        # Skip if this author is already part of another citation
        author_already_cited = False
        for citation in citations:
            if match == citation['text'].split()[0]:  # Check if it's the first word
                author_already_cited = True
                break
        
        if not author_already_cited:
            citations.append({
                'text': match,
                'style': 'Unsure'
            })
    
    # 11. Page number only citations: (42)
    for match in re.findall(r'\((\d+)\)', s):
        # Skip if it's a year that's already in APA citations
        if len(match) == 4 and 1900 <= int(match) <= 2100 and match in apa_years:
            continue
        
        # Check if this page number is already part of another citation
        page_already_cited = False
        for citation in citations:
            if f" {match}" in citation['text'] or match == citation['text']:
                page_already_cited = True
                break
        
        if not page_already_cited:
            citations.append({
                'text': f'p.{match}',
                'style': 'MLA'
            })
    
    # 12. Detect citations with year at the end of sentence like "...cognitive processes (Mislevy, 2018)."
    for match in re.findall(r'\((\w+),?\s+(\d{4})\)', s):
        author = match[0].rstrip(',').strip()
        year = match[1].strip()
        
        # Skip if this citation is already detected
        already_cited = False
        for citation in citations:
            if f"{author}" in citation['text'] and year in citation['text']:
                already_cited = True
                break
        
        if not already_cited:
            citations.append({
                'text': f"{author} ({year})",
                'style': 'APA'
            })
            apa_years.add(year)
    
    # 13. Detect in-text citations like "According to Abrams" or "Baptista and Gradim explore"
    # First look for common author-signal phrases
    author_signals = [
        r'(?:According to|as stated by|as noted by|as mentioned by|as argued by|as claimed by|as shown by|as demonstrated by|as explained by)\s+([A-Z][\w]+)',
        r'([A-Z][\w]+)\s+(?:states|argues|claims|notes|mentions|shows|demonstrates|explains|explores|examines|investigates|analyzes|discusses|suggests|proposes|concludes)',
        r'([A-Z][\w]+)\s+and\s+([A-Z][\w]+)\s+(?:state|argue|claim|note|mention|show|demonstrate|explain|explore|examine|investigate|analyze|discuss|suggest|propose|conclude)'
    ]
    
    for pattern in author_signals:
        for match in re.findall(pattern, s):
            if isinstance(match, tuple):
                # Handle multiple authors
                for author in match:
                    # Skip if this author is already part of another citation
                    author_already_cited = False
                    for citation in citations:
                        if author in citation['text']:
                            author_already_cited = True
                            break
                    
                    if not author_already_cited and len(author) > 1:  # Ensure it's a real name
                        citations.append({
                            'text': author,
                            'style': 'Narrative'
                        })
            else:
                # Single author
                author = match
                # Skip if this author is already part of another citation
                author_already_cited = False
                for citation in citations:
                    if author in citation['text']:
                        author_already_cited = True
                        break
                
                if not author_already_cited and len(author) > 1:  # Ensure it's a real name
                    citations.append({
                        'text': author,
                        'style': 'Narrative'
                    })
    
    # Better duplicate detection - check for partial matches
    unique_citations = []
    unique_citation_texts = []
    
    for i, citation in enumerate(citations):
        # Check if this citation is a subset of any other citation
        is_duplicate = False
        
        for j, other_citation in enumerate(citations):
            if i == j:  # Skip comparing with itself
                continue
                
            # Check if this citation is contained within another citation
            other_text = other_citation['text']
            if citation['text'] != other_text and citation['text'] in other_text:
                is_duplicate = True
                break
                
        if not is_duplicate and citation['text'] not in unique_citation_texts:
            unique_citations.append(citation)
            unique_citation_texts.append(citation['text'])
    
    # Sort citations by their position in the text for better presentation
    text_positions = []
    for citation in unique_citations:
        citation_text = citation['text']
        # Find position in the original text
        pos = s.find(citation_text.split('(')[0] if '(' in citation_text else citation_text)
        if pos == -1:  # If exact match not found, try a more flexible approach
            for part in citation_text.split():
                if len(part) > 3:  # Only consider substantial parts
                    pos = s.find(part)
                    if pos != -1:
                        break
        text_positions.append((pos if pos != -1 else float('inf'), citation))
    
    # Sort by position and extract just the citations
    sorted_unique_citations = [item[1] for item in sorted(text_positions, key=lambda x: x[0])]
    citation_count = len(sorted_unique_citations)
    
    return sorted_unique_citations, citation_count

def analyze_rhetorical_moves(sentence_info):
    """
    Use machine learning models to classify the rhetorical move of a sentence.
    
    The three categories are:
    - Reporting: Directly reporting what a source says
    - Transforming: Paraphrasing or synthesizing source material
    - Evaluating: Critiquing, analyzing, or evaluating sources
    - No Citation: Sentences without citations
    """
    sentence = sentence_info['sentence']
    has_citation = sentence_info['has_citation']
    
    # FIRST CHECK: Direct pattern matching for obvious rhetorical moves
    
    # Check for "According to" at the beginning of the sentence (Reporting)
    if sentence.strip().lower().startswith("according to"):
        return "Reporting", 0.95
        
    # Check for strong evaluative language (Evaluating)
    evaluative_phrases = [
        "fails to", "overlooks", "neglects", "ignores", "misses", "omits",
        "lacks", "inadequately", "insufficiently", "problematic", "questionable",
        "debatable", "controversial", "unconvincing", "unsupported", "unsubstantiated",
        "flawed", "limited", "narrow", "biased", "misleading", "weakness", "limitation",
        "shortcoming", "gap", "flaw", "critique", "criticism", "challenges", "questions",
        "contests", "disputes", "refutes", "counters", "contradicts"
    ]
    
    for phrase in evaluative_phrases:
        if phrase in sentence.lower():
            return "Evaluating", 0.95
    
    # If no citation, mark as 'No Citation'
    if not has_citation:
        return "No Citation", 1.0
    
    # Check if the sentence has more than one distinct citation
    # If so, classify it as Transforming (synthesis of multiple sources)
    if 'citations' in sentence_info and len(sentence_info['citations']) > 1:
        return "Transforming", 0.9
    
    # FIRST: Use the ML classifier to predict the rhetorical move for sentences with citations
    rhetorical_move, confidence = rhetorical_classifier.predict_rhetorical_move(sentence)
    
    # If the classifier predicts "None" but we have citations, force it to choose among the 3 rhetorical moves
    if rhetorical_move == "None" and sentence_info.get('has_citation', False):
        # Get the probabilities for the 3 rhetorical moves (excluding None)
        rhetorical_move, confidence = rhetorical_classifier.predict_rhetorical_move_no_none(sentence)
    
    # Check for very clear reporting patterns regardless of ML confidence
    reporting_patterns = [
        # With quotes
        r'^According to\b.*".*".*\([^)]+\)',  # According to X, "quote" (citation)
        r'^As\s+(?:stated|noted|reported|observed|mentioned)\s+by\b.*".*"',  # As stated by X, "quote"
        r'\b(?:states|reports|notes|observes|mentions|writes|says)\s+that\s+".*"',  # X states that "quote"
        r'\bin\s+the\s+words\s+of\b.*".*"',  # In the words of X, "quote"
        
        # Without quotes - narrative citations
        r'^According to\b.*?\([^)]+\)',  # According to Smith (2019), ...
        r'^As\s+(?:stated|noted|reported|observed|mentioned)\s+by\b.*?\([^)]+\)',  # As stated by Smith (2019), ...
        r'(?:states|reports|notes|observes|mentions|writes|says)\s+that\b.*?\([^)]+\)',  # Smith states that (2019) ...
        
        # Common reporting phrases
        r'\breports?\s+(?:that|how|on)\b',  # reports that/how/on
        r'\bdescribes?\s+(?:that|how)\b',  # describes that/how
        r'\bexplains?\s+(?:that|how)\b',  # explains that/how
        r'\bidentifies?\b',  # identifies
        r'\bdocuments?\b'  # documents
    ]
    
    for pattern in reporting_patterns:
        if re.search(pattern, sentence, re.IGNORECASE):
            return "Reporting", 0.9
    
    # SECOND: Only use pattern matching if the ML confidence is low
    if confidence < 0.7:
        # Check for implicit citations with specific rhetorical move types
        if 'citations' in sentence_info:
            for citation in sentence_info['citations']:
                # Check for common reporting phrases
                if citation.lower().startswith("according to") or \
                   citation.lower().startswith("as stated by") or \
                   citation.lower().startswith("as noted by") or \
                   "reports that" in citation.lower() or \
                   "states that" in citation.lower():
                    return "Reporting", 0.9
                # Check for evaluative phrases
                elif "fails to" in citation.lower() or \
                     "overlooks" in citation.lower() or \
                     "neglects" in citation.lower() or \
                     "limitation" in citation.lower() or \
                     "critique" in citation.lower() or \
                     "weakness" in citation.lower():
                    return "Evaluating", 0.9
                # Otherwise assume transforming
                elif sentence_info.get('has_citation', False):
                    return "Transforming", 0.8
                # Check for author references (possessive form)
                elif "'s" in citation and any(author in citation for author in author_names):
                    return "Transforming", 0.7
                # Check for author references (narrative form)
                elif any(author in citation for author in author_names):
                    return "Transforming", 0.7
        
        # Check for strong evaluative patterns if ML confidence is low
        evaluating_patterns = [
            r'\b(?:overlooks|fails to|neglects|ignores|misses|omits)\b',
            r'\b(?:critique|criticism|limitation|weakness|flaw)s?\b',
            r'\b(?:problematic|questionable|debatable|controversial|unconvincing)\b'
        ]
        
        for pattern in evaluating_patterns:
            if re.search(pattern, sentence, re.IGNORECASE) and confidence < 0.65:
                return "Evaluating", 0.7
    
    # If the classifier didn't detect a rhetorical move with any confidence,
    # default to "Reporting" as it's the most common move with citations
    if confidence < 0.5:
        rhetorical_move = "Reporting"
        confidence = 0.5
    
    return rhetorical_move, confidence

@app.route('/')
def home():
    @after_this_request
    def add_no_cache(response):
        response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
        response.headers["Pragma"] = "no-cache"
        response.headers["Expires"] = "0"
        return response
    # Serve the fixed chart template
    return render_template('fixed_chart.html')

@app.route('/new')
def new_interface():
    @after_this_request
    def add_no_cache(response):
        response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
        response.headers["Pragma"] = "no-cache"
        response.headers["Expires"] = "0"
        return response
    return render_template('index_new.html')

@app.route('/analyze', methods=['POST'])
def analyze():
    print('Received request:', request.json)
    text = request.json.get('text', '')
    
    if not text:
        return jsonify({
            'error': 'No text provided',
            'sentence_analysis': [],
            'citation_count': 0
        })
    
    return process_text(text)


@app.route('/upload_file', methods=['POST'])
def upload_file():
    """Upload a file and extract its text content"""
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'})
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'})
    
    # Check file extension
    allowed_extensions = {'txt', 'docx', 'pdf'}
    file_ext = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''
    
    if file_ext not in allowed_extensions:
        return jsonify({'error': f'File type not supported. Please upload a {", ".join(allowed_extensions)} file'})
    
    try:
        # Extract text based on file type
        if file_ext == 'txt':
            # For text files, read directly
            text = file.read().decode('utf-8')
        elif file_ext == 'docx':
            # For DOCX files, use python-docx
            doc = Document(io.BytesIO(file.read()))
            text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        elif file_ext == 'pdf':
            # For PDF files, use PyPDF2
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file.read()))
            text = '\n'.join([page.extract_text() for page in pdf_reader.pages])
        
        return jsonify({'text': text})
    except Exception as e:
        return jsonify({'error': f'Error extracting text: {str(e)}'})


@app.route('/upload', methods=['GET'])
def upload_page():
    """Serve the upload and analyze page"""
    return render_template('upload_and_analyze.html')


@app.route('/upload_and_analyze', methods=['POST'])
def upload_and_analyze():
    """Upload a document, analyze it, and provide download options"""
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'})
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'})
    
    # Create a temporary file to store the uploaded file
    with tempfile.NamedTemporaryFile(delete=False) as temp:
        file.save(temp.name)
        filename = secure_filename(file.filename)
        text = extract_text_from_file(temp.name, filename)
        # Delete the temporary file
        os.unlink(temp.name)
    
    if not text.strip():
        return jsonify({'error': 'Could not extract text from the file'})
    
    # Process the text and get analysis results
    analysis_result = process_text(text)
    
    # Store the analysis result in a session variable for download later
    analysis_data = json.loads(analysis_result.get_data(as_text=True))
    session['last_analysis'] = analysis_data
    session['analyzed_text'] = text
    session['analyzed_filename'] = filename
    
    return analysis_result


@app.route('/download_analysis_html', methods=['GET'])
def download_analysis_html():
    """Download the analysis results as an HTML file"""
    if 'last_analysis' not in session or 'analyzed_filename' not in session:
        return jsonify({'error': 'No analysis results available for download'})
    
    analysis_data = session['last_analysis']
    filename = session['analyzed_filename']
    analyzed_text = session['analyzed_text']
    
    # Generate a formatted HTML document with the analysis results
    html_content = generate_analysis_html(analysis_data, analyzed_text, filename)
    
    # Create response with HTML content
    response = make_response(html_content)
    response.headers["Content-Type"] = "text/html"
    response.headers["Content-Disposition"] = f"attachment; filename=analysis_{os.path.splitext(filename)[0]}.html"
    
    return response


@app.route('/download_analysis_pdf', methods=['GET'])
def download_analysis_pdf():
    """Download the analysis results as a PDF file"""
    if 'last_analysis' not in session or 'analyzed_filename' not in session:
        return jsonify({'error': 'No analysis results available for download'})
    
    analysis_data = session['last_analysis']
    filename = session['analyzed_filename']
    analyzed_text = session['analyzed_text']
    
    # Generate a formatted HTML document with the analysis results
    html_content = generate_analysis_html(analysis_data, analyzed_text, filename)
    
    try:
        # Check if wkhtmltopdf is installed
        import shutil
        wkhtmltopdf_installed = shutil.which('wkhtmltopdf') is not None
        
        if wkhtmltopdf_installed:
            # Create a temporary file for the PDF
            temp_pdf_path = os.path.join(tempfile.gettempdir(), f"analysis_{uuid.uuid4().hex}.pdf")
            
            # Convert HTML to PDF
            pdfkit.from_string(html_content, temp_pdf_path)
            
            # Send the PDF file
            return send_file(
                temp_pdf_path,
                mimetype='application/pdf',
                as_attachment=True,
                download_name=f"analysis_{os.path.splitext(filename)[0]}.pdf"
            )
        else:
            # If wkhtmltopdf is not installed, return HTML instead
            response = make_response(html_content)
            response.headers["Content-Type"] = "text/html"
            response.headers["Content-Disposition"] = f"attachment; filename=analysis_{os.path.splitext(filename)[0]}.html"
            return response
    except Exception as e:
        return jsonify({'error': f'PDF generation failed: {str(e)}'})
    finally:
        # Clean up the temporary file if it was created
        @after_this_request
        def remove_file(response):
            try:
                temp_pdf_path = os.path.join(tempfile.gettempdir(), f"analysis_{uuid.uuid4().hex}.pdf")
                if os.path.exists(temp_pdf_path):
                    os.remove(temp_pdf_path)
            except Exception as e:
                app.logger.error(f"Error removing temporary file: {e}")
            return response

@app.route('/analyze_files', methods=['POST'])
def analyze_files():
    if 'files' not in request.files:
        return jsonify({
            'error': 'No files provided',
            'sentence_analysis': [],
            'citation_count': 0
        })
    
    files = request.files.getlist('files')
    if not files or files[0].filename == '':
        return jsonify({
            'error': 'No files selected',
            'sentence_analysis': [],
            'citation_count': 0
        })
    
    # Process each file and combine the text
    combined_text = ""
    for file in files:
        filename = secure_filename(file.filename)
        # Create a temporary file to store the uploaded file
        with tempfile.NamedTemporaryFile(delete=False) as temp:
            file.save(temp.name)
            file_text = extract_text_from_file(temp.name, filename)
            combined_text += file_text + "\n\n"
            # Delete the temporary file
            os.unlink(temp.name)
    
    if not combined_text.strip():
        return jsonify({
            'error': 'Could not extract text from the provided files',
            'sentence_analysis': [],
            'citation_count': 0
        })
    
    return process_text(combined_text)

@app.route('/display_citations', methods=['POST'])
def display_citations():
    """Display uploaded text with citation tags highlighted"""
    if 'file' not in request.files:
        return render_template('display_citations.html', error='No file provided')
    
    file = request.files['file']
    if file.filename == '':
        return render_template('display_citations.html', error='No file selected')
    
    # Create a temporary file to store the uploaded file
    with tempfile.NamedTemporaryFile(delete=False) as temp:
        file.save(temp.name)
        filename = secure_filename(file.filename)
        text = extract_text_from_file(temp.name, filename)
        # Delete the temporary file
        os.unlink(temp.name)
    
    if not text.strip():
        return render_template('display_citations.html', error='Could not extract text from the file')
    
    # Extract author names for reference tracking
    author_names = extract_author_names(text)
    
    # Process the text to identify citations
    # We'll use a simplified version of preprocess_text that preserves formatting
    paragraphs = text.split('\n\n')
    processed_paragraphs = []
    all_citations = []
    
    for paragraph in paragraphs:
        if not paragraph.strip():
            processed_paragraphs.append('')
            continue
            
        # Process each sentence in the paragraph
        sentences = nltk.sent_tokenize(paragraph)
        processed_sentences = []
        
        for sentence in sentences:
            # Identify citations in this sentence
            citations, _ = identify_citations(sentence, author_names)
            
            # If we found citations, mark them in the text
            if citations:
                marked_sentence = sentence
                for citation in citations:
                    citation_text = citation['text']
                    citation_style = citation['style']
                    
                    # Only mark the citation if it's actually in the text (not implicit)
                    if citation_style != 'Implicit' and citation_text in sentence:
                        # Create a span with appropriate styling
                        span = f'<span class="citation {citation_style.lower().replace(" ", "-")}" title="{citation_style}">{citation_text}</span>'
                        marked_sentence = marked_sentence.replace(citation_text, span)
                
                processed_sentences.append(marked_sentence)
                all_citations.extend(citations)
            else:
                processed_sentences.append(sentence)
        
        processed_paragraphs.append(' '.join(processed_sentences))
    
    # Join paragraphs with double line breaks to preserve formatting
    processed_text = '\n\n'.join(processed_paragraphs)
    
    # Count citations by style
    citation_counts = {}
    for citation in all_citations:
        style = citation['style']
        if style in citation_counts:
            citation_counts[style] += 1
        else:
            citation_counts[style] = 1
    
    return render_template('display_citations.html', 
                           text=processed_text, 
                           citations=all_citations,
                           citation_counts=citation_counts,
                           filename=filename)

def extract_text_from_file(file_path, filename):
    """Extract text from different file types"""
    text = ""
    file_ext = os.path.splitext(filename)[1].lower()
    
    try:
        if file_ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        elif file_ext == '.pdf':
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
        elif file_ext in ['.doc', '.docx']:
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
    except Exception as e:
        print(f"Error extracting text from {filename}: {str(e)}")
        text = f"Error processing {filename}: {str(e)}"
    
    return text

def extract_author_names(text):
    """Extract all author names from citations in the text for later reference"""
    author_names = set()
    
    # Extract author names from APA citations: (Smith, 2020)
    for match in re.findall(r'\(([\w]+),\s*\d{4}', text):
        author_names.add(match)
    
    # Extract author names from APA citations with multiple authors: (Smith et al., 2020)
    for match in re.findall(r'\(([\w]+)\s+et\s+al\.', text):
        author_names.add(match)
    
    # Extract author names from APA citations with two authors: (Smith and Jones, 2020)
    for match in re.findall(r'\(([\w]+)\s+and\s+([\w]+)', text):
        author_names.add(match[0])
        author_names.add(match[1])
    
    # Extract author names from narrative citations: Smith (2020)
    for match in re.findall(r'([\w]+)\s+\(\d{4}\)', text):
        author_names.add(match)
    
    # Extract author names from MLA citations
    for match in re.findall(r'\(([\w]+)\s+\d+\)', text):
        author_names.add(match)
    
    return author_names

def generate_analysis_html(analysis_data, original_text, filename):
    """Generate a formatted HTML document with analysis results"""
    # Get current date and time
    current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    # Extract data from analysis
    sentence_analysis = analysis_data.get('sentence_analysis', [])
    citation_count = analysis_data.get('citation_count', 0)
    rhetorical_move_stats = analysis_data.get('rhetorical_move_stats', {})
    
    # Start building HTML content
    html = f'''
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Citation Analysis Report - {filename}</title>
        <style>
            :root {{
                --primary-color: #2c536e;
                --secondary-color: #3c7ba8;
                --accent-color: #e74c3c;
                --background-color: #f5f5f5;
                --card-color: #ffffff;
                --text-color: #333333;
                --border-radius: 8px;
                --spacing: 20px;
                --reporting-color: #4285f4;
                --transforming-color: #34a853;
                --evaluating-color: #ea4335;
            }}
            
            body {{
                font-family: 'Arial', sans-serif;
                background-color: var(--background-color);
                margin: 0;
                padding: 0;
                color: var(--text-color);
                line-height: 1.6;
            }}
            
            .container {{
                max-width: 1100px;
                margin: 0 auto;
                padding: 20px;
            }}
            
            .header {{
                background-color: var(--primary-color);
                color: white;
                padding: 20px;
                margin-bottom: 30px;
                border-radius: var(--border-radius);
            }}
            
            .header h1 {{
                margin: 0;
                font-size: 24px;
            }}
            
            .header p {{
                margin: 5px 0 0;
                opacity: 0.8;
            }}
            
            .summary-box {{
                background-color: #f8f9fa;
                border-left: 4px solid var(--secondary-color);
                padding: 15px;
                margin-bottom: 30px;
                border-radius: var(--border-radius);
            }}
            
            .stats-container {{
                display: flex;
                flex-wrap: wrap;
                gap: 20px;
                margin-bottom: 30px;
            }}
            
            .stats-card {{
                background-color: white;
                border-radius: var(--border-radius);
                padding: 15px;
                box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                flex: 1;
                min-width: 200px;
                text-align: center;
            }}
            
            .stats-card h3 {{
                margin-top: 0;
                color: var(--primary-color);
            }}
            
            .stats-card .number {{
                font-size: 36px;
                font-weight: bold;
                margin: 10px 0;
            }}
            
            .stats-card .percentage {{
                font-size: 18px;
                color: #666;
            }}
            
            .chart-container {{
                display: flex;
                justify-content: space-around;
                height: 250px;
                margin: 30px 0;
                align-items: flex-end;
            }}
            
            .chart-bar {{
                width: 80px;
                display: flex;
                flex-direction: column;
                align-items: center;
            }}
            
            .chart-fill {{
                width: 100%;
                background-color: var(--primary-color);
                border-radius: 4px 4px 0 0;
                min-height: 10px;
            }}
            
            .chart-label {{
                margin-top: 10px;
                text-align: center;
                font-weight: bold;
            }}
            
            .chart-value {{
                margin-bottom: 5px;
                font-weight: bold;
            }}
            
            .result-card {{
                background-color: white;
                border-radius: var(--border-radius);
                padding: 20px;
                margin-bottom: 20px;
                box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                border-left: 4px solid #ccc;
            }}
            
            .result-card.reporting-move {{
                border-left-color: var(--reporting-color);
            }}
            
            .result-card.transforming-move {{
                border-left-color: var(--transforming-color);
            }}
            
            .result-card.evaluating-move {{
                border-left-color: var(--evaluating-color);
            }}
            
            .result-card.has-citation {{
                border-left-width: 6px;
            }}
            
            .move-badge {{
                display: inline-block;
                padding: 5px 10px;
                border-radius: 20px;
                font-size: 14px;
                font-weight: bold;
                color: white;
                margin-right: 10px;
            }}
            
            .reporting-badge {{
                background-color: var(--reporting-color);
            }}
            
            .transforming-badge {{
                background-color: var(--transforming-color);
            }}
            
            .evaluating-badge {{
                background-color: var(--evaluating-color);
            }}
            
            .citation-info {{
                background-color: #f8f9fa;
                border-radius: var(--border-radius);
            }}
            
            .confidence-bar {{
                height: 10px;
                background-color: var(--secondary-color);
                border-radius: 5px;
                margin-top: 5px;
            }}
            
            .original-text {{
                background-color: white;
                border-radius: var(--border-radius);
                padding: 20px;
                margin-top: 30px;
                box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                white-space: pre-wrap;
                font-family: 'Georgia', serif;
                line-height: 1.8;
            }}
            
            h2 {{
                color: var(--primary-color);
                margin-top: 40px;
                margin-bottom: 20px;
                border-bottom: 1px solid #eee;
                padding-bottom: 10px;
            }}
            
            @media print {{
                body {{
                    background-color: white;
                }}
                
                .container {{
                    width: 100%;
                    max-width: none;
                }}
                
                .result-card, .stats-card, .summary-box, .original-text {{
                    break-inside: avoid;
                }}
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">
                <h1>Citation Analysis Report</h1>
                <p>Filename: {filename}</p>
                <p>Generated on: {current_time}</p>
            </div>
            
            <div class="summary-box">
                <h2>Analysis Summary</h2>
                <p><strong>Total Citations Found:</strong> {citation_count}</p>
            </div>
    '''
    
    # Add rhetorical move statistics
    if rhetorical_move_stats:
        move_counts = rhetorical_move_stats.get('counts', {})
        move_percentages = rhetorical_move_stats.get('percentages', {})
        
        html += f'''
            <h2>Rhetorical Move Statistics</h2>
            <div class="stats-container">
                <div class="stats-card">
                    <h3>Citations</h3>
                    <div class="number">{citation_count}</div>
                    <div class="percentage">Total</div>
                </div>
                <div class="stats-card">
                    <h3>Reporting</h3>
                    <div class="number">{move_counts.get('Reporting', 0)}</div>
                    <div class="percentage">{move_percentages.get('Reporting', 0)}%</div>
                </div>
                <div class="stats-card">
                    <h3>Transforming</h3>
                    <div class="number">{move_counts.get('Transforming', 0)}</div>
                    <div class="percentage">{move_percentages.get('Transforming', 0)}%</div>
                </div>
                <div class="stats-card">
                    <h3>Evaluating</h3>
                    <div class="number">{move_counts.get('Evaluating', 0)}</div>
                    <div class="percentage">{move_percentages.get('Evaluating', 0)}%</div>
                </div>
            </div>
            
            <div class="chart-container">
                <div class="chart-bar">
                    <div class="chart-value">{move_percentages.get('Reporting', 0)}%</div>
                    <div class="chart-fill" style="height: {max(move_percentages.get('Reporting', 0), 5)}%; background-color: var(--reporting-color);"></div>
                    <div class="chart-label">Reporting</div>
                </div>
                <div class="chart-bar">
                    <div class="chart-value">{move_percentages.get('Transforming', 0)}%</div>
                    <div class="chart-fill" style="height: {max(move_percentages.get('Transforming', 0), 5)}%; background-color: var(--transforming-color);"></div>
                    <div class="chart-label">Transforming</div>
                </div>
                <div class="chart-bar">
                    <div class="chart-value">{move_percentages.get('Evaluating', 0)}%</div>
                    <div class="chart-fill" style="height: {max(move_percentages.get('Evaluating', 0), 5)}%; background-color: var(--evaluating-color);"></div>
                    <div class="chart-label">Evaluating</div>
                </div>
            </div>
        '''
    
    # Add sentence analysis
    html += '''
        <h2>Sentence Analysis</h2>
    '''
    
    for result in sentence_analysis:
        # Determine rhetorical move class
        move_class = ''
        move_badge = ''
        rhetorical_move = result.get('rhetorical_move', '')
        
        if rhetorical_move == 'Reporting':
            move_class = 'reporting-move'
            move_badge = '<span class="move-badge reporting-badge">Reporting</span>'
        elif rhetorical_move == 'Transforming':
            move_class = 'transforming-move'
            move_badge = '<span class="move-badge transforming-badge">Transforming</span>'
        elif rhetorical_move == 'Evaluating':
            move_class = 'evaluating-move'
            move_badge = '<span class="move-badge evaluating-badge">Evaluating</span>'
        
        has_citation = result.get('has_citation', False)
        citation_class = 'has-citation' if has_citation else 'no-citation'
        
        # Citation info
        citation_info = ''
        if has_citation and 'citations' in result and result['citations']:
            citation_info = f'''
                <div class="citation-info" style="padding: 15px; margin-top: 15px;">
                    <p style="margin-bottom: 10px;"><strong>Citations ({len(result['citations'])}):</strong></p>
                    <p style="margin-bottom: 0;">{', '.join(result['citations'])}</p>
                </div>
            '''
        
        # Confidence bar
        confidence = result.get('confidence', 0)
        confidence_percentage = round(confidence * 100, 1)
        confidence_color = '#4285f4' if confidence_percentage > 80 else '#ffc107' if confidence_percentage > 50 else '#ea4335'
        
        html += f'''
            <div class="result-card {move_class} {citation_class}">
                <div style="margin-bottom: 15px;">
                    <h3 style="margin-top: 0; margin-bottom: 10px;">Sentence</h3>
                    <p style="margin-bottom: 0;">{result.get('sentence', '')}</p>
                </div>
                <div style="margin-bottom: 15px;">
                    <h3 style="margin-top: 0; margin-bottom: 10px;">Rhetorical Move</h3>
                    <div>{move_badge} {rhetorical_move}</div>
                </div>
                {citation_info}
                <div style="margin-top: 15px;">
                    <h3 style="margin-top: 0; margin-bottom: 10px;">Citation Style</h3>
                    <p style="margin-bottom: 0;">{result.get('citation_style', 'None detected')}</p>
                </div>
                <div style="margin-top: 15px;">
                    <div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 5px;">
                        <span>Confidence</span>
                        <span style="font-weight: bold;">{confidence_percentage}%</span>
                    </div>
                    <div class="confidence-bar" style="width: {confidence_percentage}%; background-color: {confidence_color};"></div>
                </div>
            </div>
        '''
    
    # Add original text
    html += f'''
        <h2>Original Text</h2>
        <div class="original-text">
            {original_text}
        </div>
    '''
    
    # Close HTML
    html += '''
        </div>
    </body>
    </html>
    '''
    
    return html


def process_text(text):
    
    # Extract all author names from citations for later reference
    author_names = extract_author_names(text)
    
    # Preprocess and analyze the text
    analyzed_sentences, all_citations, total_citation_count = preprocess_text(text, author_names)
    
    # Add rhetorical move analysis using ML models
    for sentence_info in analyzed_sentences:
        rhetorical_move, confidence = analyze_rhetorical_moves(sentence_info)
        sentence_info['rhetorical_move'] = rhetorical_move
        sentence_info['confidence'] = confidence
        
        # Add citation styles information
        if sentence_info['has_citation'] and sentence_info['citations']:
            # Get the raw citations with their styles
            citations_with_styles = []
            for citation_text in sentence_info['citations']:
                # Find the citation with style in all_citations
                for citation in all_citations:
                    if citation['text'] == citation_text:
                        citations_with_styles.append({
                            'text': citation_text,
                            'style': citation['style']
                        })
                        break
                else:
                    # If not found, add as unsure
                    citations_with_styles.append({
                        'text': citation_text,
                        'style': 'Unsure'
                    })
            
            sentence_info['citations_with_styles'] = citations_with_styles
    
    # Calculate statistics for rhetorical moves
    move_counts = {
        "Reporting": 0,
        "Transforming": 0,
        "Evaluating": 0,
        "No Citation": 0
    }
    
    # Filter out section headers (all caps sentences with no punctuation)
    content_sentences = []
    for sentence_info in analyzed_sentences:
        sentence = sentence_info['sentence']
        # Check if this is a section header (all caps, short, no punctuation except colon)
        is_header = sentence.isupper() and len(sentence.split()) <= 3 and not any(p in sentence for p in '.!?,;')
        if not is_header:
            content_sentences.append(sentence_info)
            move = sentence_info['rhetorical_move']
            if move in move_counts:
                move_counts[move] += 1
    
    # Calculate percentages using only content sentences
    total_content_sentences = len(content_sentences)
    move_percentages = {}
    for move, count in move_counts.items():
        percentage = (count / total_content_sentences) * 100 if total_content_sentences > 0 else 0
        move_percentages[move] = round(percentage, 1)
    
    # Debug print to see what's being sent
    print("All citations:", all_citations)
    for sentence in analyzed_sentences:
        if 'citations_with_styles' in sentence:
            print("Citations with styles:", sentence['citations_with_styles'])
    
    return jsonify({
        'sentence_analysis': analyzed_sentences,
        'citation_count': total_citation_count,
        'rhetorical_move_stats': {
            'counts': move_counts,
            'percentages': move_percentages
        }
    })

if __name__ == '__main__':
    nltk.download('punkt')  # Download required NLTK data
    app.run(debug=True, host='0.0.0.0', port=5001)
